"""
ExtractWorker - Text extraction from PDF, DOCX, and XLSX files.

Extracts text content from common document formats using CPU-based libraries.
Part of the cpu-extract worker pool for document processing.
"""

import asyncio
import logging
import os
from pathlib import Path
from typing import Any, Dict, Optional

from arkham_frame.workers.base import BaseWorker

logger = logging.getLogger(__name__)


class ExtractWorker(BaseWorker):
    """
    Worker for extracting text from documents.

    Supports PDF, DOCX, and XLSX file formats. Handles various edge cases
    including missing files, corrupted documents, password-protected files,
    and missing dependencies.

    Uses the cpu-extract pool for document text extraction tasks.
    """

    pool = "cpu-extract"
    name = "ExtractWorker"

    # Configuration
    poll_interval = 1.0
    heartbeat_interval = 10.0
    idle_timeout = 300.0  # 5 minutes
    job_timeout = 120.0  # 2 minutes for large files
    max_retries = 2

    def __init__(self, *args, **kwargs):
        """Initialize worker and check for required dependencies."""
        super().__init__(*args, **kwargs)
        self._check_dependencies()

    def _check_dependencies(self):
        """Check which extraction libraries are available."""
        self._has_pypdf = False
        self._has_docx = False
        self._has_openpyxl = False

        try:
            import pypdf

            self._has_pypdf = True
        except ImportError:
            logger.warning("pypdf not installed - PDF extraction unavailable")

        try:
            import docx

            self._has_docx = True
        except ImportError:
            logger.warning("python-docx not installed - DOCX extraction unavailable")

        try:
            import openpyxl

            self._has_openpyxl = True
        except ImportError:
            logger.warning("openpyxl not installed - XLSX extraction unavailable")

        if not any([self._has_pypdf, self._has_docx, self._has_openpyxl]):
            logger.error("No extraction libraries available! Install pypdf, python-docx, and/or openpyxl")

    def _resolve_path(self, file_path: str) -> Path:
        """
        Resolve file path using DATA_SILO_PATH for Docker/portable deployments.

        Args:
            file_path: Path from payload (may be relative or absolute)

        Returns:
            Resolved absolute Path
        """
        if not os.path.isabs(file_path):
            data_silo = os.environ.get("DATA_SILO_PATH", ".")
            return Path(data_silo) / file_path
        return Path(file_path)

    async def process_job(self, job_id: str, payload: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract text from a document.

        Payload:
            file_path: Path to the file to extract from (required)
            file_type: File type - "pdf", "docx", or "xlsx" (required)

        Returns:
            dict with:
                success: bool - Whether extraction succeeded
                text: str - Extracted text content
                pages: int - Number of pages/sheets processed
                error: str - Error message if success=False
                file_path: str - Original file path
                file_type: str - File type processed

        Raises:
            ValueError: If required parameters are missing or invalid
            FileNotFoundError: If file doesn't exist
            Exception: For other extraction errors
        """
        # Validate payload
        file_path = payload.get("file_path")
        file_type = payload.get("file_type", "").lower()

        if not file_path:
            raise ValueError("Missing required parameter: file_path")

        # Resolve relative path using DATA_SILO_PATH
        file_path = str(self._resolve_path(file_path))

        # Auto-detect file_type from extension if not provided
        if not file_type:
            path = Path(file_path)
            ext = path.suffix.lower()
            ext_to_type = {
                ".pdf": "pdf",
                ".docx": "docx",
                ".doc": "docx",  # Old Word format (may not work)
                ".xlsx": "xlsx",
                ".xls": "xlsx",  # Old Excel format (may not work)
                ".csv": "csv",
                ".tsv": "csv",  # Tab-separated values
                ".txt": "txt",
                ".text": "txt",
                ".md": "txt",
                ".log": "txt",
                ".eml": "eml",
                ".emlx": "eml",  # Apple Mail format
            }
            file_type = ext_to_type.get(ext, "")
            if file_type:
                logger.debug(f"Auto-detected file_type: {file_type} from extension {ext}")

        if not file_type:
            raise ValueError(
                "Could not determine file type. "
                "Provide file_type parameter or use a supported extension: "
                "pdf, docx, xlsx, csv, txt, eml"
            )

        if file_type not in ["pdf", "docx", "xlsx", "csv", "txt", "eml"]:
            raise ValueError(f"Unsupported file_type: {file_type}. Supported types: pdf, docx, xlsx, csv, txt, eml")

        # Check file exists
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if not path.is_file():
            raise ValueError(f"Path is not a file: {file_path}")

        logger.info(f"ExtractWorker processing {file_type.upper()} file: {path.name} (job {job_id})")

        # Dispatch to appropriate extractor
        try:
            if file_type == "pdf":
                result = await self._extract_pdf(path)
            elif file_type == "docx":
                result = await self._extract_docx(path)
            elif file_type == "xlsx":
                result = await self._extract_xlsx(path)
            elif file_type == "txt":
                result = await self._extract_text(path)
            elif file_type == "csv":
                result = await self._extract_csv(path)
            elif file_type == "eml":
                result = await self._extract_eml(path)
            else:
                # Should never reach here due to earlier validation
                raise ValueError(f"Unsupported file type: {file_type}")

            # Add metadata
            result["file_path"] = str(file_path)
            result["file_type"] = file_type
            result["success"] = True

            logger.info(
                f"ExtractWorker completed {file_type.upper()}: "
                f"{result.get('pages', 0)} pages, "
                f"{len(result.get('text', ''))} chars"
            )

            return result

        except Exception as e:
            error_msg = f"Extraction failed for {file_type.upper()}: {str(e)}"
            logger.error(f"ExtractWorker error (job {job_id}): {error_msg}")

            return {
                "success": False,
                "text": "",
                "pages": 0,
                "error": error_msg,
                "file_path": str(file_path),
                "file_type": file_type,
            }

    async def _extract_pdf(self, path: Path) -> Dict[str, Any]:
        """
        Extract text and metadata from PDF file.

        Args:
            path: Path to PDF file

        Returns:
            dict with text, page count, and document_metadata

        Raises:
            ImportError: If pypdf is not installed
            Exception: For PDF reading errors
        """
        if not self._has_pypdf:
            raise ImportError("pypdf library not installed. Install with: pip install pypdf")

        from pypdf import PdfReader

        # Run in executor to avoid blocking
        def extract():
            try:
                reader = PdfReader(str(path))

                # Check for encryption
                if reader.is_encrypted:
                    raise ValueError("PDF is password-protected. Encrypted PDFs are not supported.")

                # Extract PDF metadata (author, title, creator, etc.)
                document_metadata = {}
                if reader.metadata:
                    meta = reader.metadata
                    # Standard PDF metadata fields
                    if meta.author:
                        document_metadata["author"] = str(meta.author)
                    if meta.title:
                        document_metadata["title"] = str(meta.title)
                    if meta.subject:
                        document_metadata["subject"] = str(meta.subject)
                    if meta.creator:
                        document_metadata["creator"] = str(meta.creator)
                    if meta.producer:
                        document_metadata["producer"] = str(meta.producer)
                    if meta.creation_date:
                        document_metadata["creation_date"] = str(meta.creation_date)
                    if meta.modification_date:
                        document_metadata["modification_date"] = str(meta.modification_date)
                    # Keywords (may be comma-separated string)
                    if hasattr(meta, "keywords") and meta.keywords:
                        document_metadata["keywords"] = str(meta.keywords)

                pages = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)

                full_text = "\n\n".join(pages)

                return {
                    "text": full_text,
                    "pages": len(reader.pages),
                    "document_metadata": document_metadata,
                }

            except Exception as e:
                # Add context to error
                raise Exception(f"PDF reading error: {str(e)}")

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract)

    async def _extract_docx(self, path: Path) -> Dict[str, Any]:
        """
        Extract text and metadata from DOCX file.

        Args:
            path: Path to DOCX file

        Returns:
            dict with text, paragraph count, and document_metadata

        Raises:
            ImportError: If python-docx is not installed
            Exception: For DOCX reading errors
        """
        if not self._has_docx:
            raise ImportError("python-docx library not installed. Install with: pip install python-docx")

        from docx import Document

        # Run in executor to avoid blocking
        def extract():
            try:
                doc = Document(str(path))

                # Extract document metadata from core properties
                document_metadata = {}
                if doc.core_properties:
                    props = doc.core_properties
                    if props.author:
                        document_metadata["author"] = str(props.author)
                    if props.title:
                        document_metadata["title"] = str(props.title)
                    if props.subject:
                        document_metadata["subject"] = str(props.subject)
                    if props.keywords:
                        document_metadata["keywords"] = str(props.keywords)
                    if props.category:
                        document_metadata["category"] = str(props.category)
                    if props.comments:
                        document_metadata["comments"] = str(props.comments)
                    if props.last_modified_by:
                        document_metadata["last_modified_by"] = str(props.last_modified_by)
                    if props.created:
                        document_metadata["creation_date"] = str(props.created)
                    if props.modified:
                        document_metadata["modification_date"] = str(props.modified)
                    if props.revision:
                        document_metadata["revision"] = str(props.revision)

                # Extract paragraphs
                paragraphs = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        paragraphs.append(text)

                # Extract table content
                tables = []
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        if any(cells):  # Skip empty rows
                            tables.append(" | ".join(cells))

                # Combine all text
                all_text = []
                all_text.extend(paragraphs)
                if tables:
                    all_text.append("\n--- Tables ---\n")
                    all_text.extend(tables)

                full_text = "\n".join(all_text)

                return {
                    "text": full_text,
                    "pages": len(paragraphs),  # Use paragraph count as proxy
                    "document_metadata": document_metadata,
                }

            except Exception as e:
                raise Exception(f"DOCX reading error: {str(e)}")

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract)

    async def _extract_xlsx(self, path: Path) -> Dict[str, Any]:
        """
        Extract text and metadata from XLSX file.

        Args:
            path: Path to XLSX file

        Returns:
            dict with text, sheet count, and document_metadata

        Raises:
            ImportError: If openpyxl is not installed
            Exception: For XLSX reading errors
        """
        if not self._has_openpyxl:
            raise ImportError("openpyxl library not installed. Install with: pip install openpyxl")

        from openpyxl import load_workbook

        # Run in executor to avoid blocking
        def extract():
            try:
                # Load workbook (NOT read-only so we can access properties)
                wb = load_workbook(
                    str(path),
                    read_only=False,  # Need full load for properties
                    data_only=True,  # Get computed values, not formulas
                )

                # Extract workbook metadata
                document_metadata = {}
                if wb.properties:
                    props = wb.properties
                    if props.creator:
                        document_metadata["author"] = str(props.creator)
                    if props.title:
                        document_metadata["title"] = str(props.title)
                    if props.subject:
                        document_metadata["subject"] = str(props.subject)
                    if props.description:
                        document_metadata["description"] = str(props.description)
                    if props.keywords:
                        document_metadata["keywords"] = str(props.keywords)
                    if props.category:
                        document_metadata["category"] = str(props.category)
                    if props.lastModifiedBy:
                        document_metadata["last_modified_by"] = str(props.lastModifiedBy)
                    if props.created:
                        document_metadata["creation_date"] = str(props.created)
                    if props.modified:
                        document_metadata["modification_date"] = str(props.modified)
                    if props.company:
                        document_metadata["company"] = str(props.company)

                sheets = []

                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]

                    sheet_text = [f"--- Sheet: {sheet_name} ---"]

                    # Extract cell values row by row
                    for row in sheet.iter_rows(values_only=True):
                        # Convert to strings and filter out None
                        cells = [str(cell) for cell in row if cell is not None]
                        if cells:
                            sheet_text.append(" | ".join(cells))

                    sheets.append("\n".join(sheet_text))

                full_text = "\n\n".join(sheets)

                return {
                    "text": full_text,
                    "pages": len(wb.sheetnames),
                    "document_metadata": document_metadata,
                }

            except Exception as e:
                raise Exception(f"XLSX reading error: {str(e)}")

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract)

    async def _extract_text(self, path: Path) -> Dict[str, Any]:
        """
        Extract text from plain text file.

        Args:
            path: Path to text file

        Returns:
            dict with text and line count
        """

        def extract():
            try:
                # Try different encodings
                encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
                text = None

                for encoding in encodings:
                    try:
                        with open(path, "r", encoding=encoding) as f:
                            text = f.read()
                        break
                    except UnicodeDecodeError:
                        continue

                if text is None:
                    # Fallback: read as bytes and decode with errors ignored
                    with open(path, "rb") as f:
                        text = f.read().decode("utf-8", errors="replace")

                lines = text.count("\n") + 1

                return {
                    "text": text,
                    "pages": lines,  # Use line count as proxy for pages
                }

            except Exception as e:
                raise Exception(f"Text file reading error: {str(e)}")

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract)

    async def _extract_csv(self, path: Path) -> Dict[str, Any]:
        """
        Extract text from CSV or TSV file.

        Converts tabular data to readable text format with headers and rows.

        Args:
            path: Path to CSV/TSV file

        Returns:
            dict with text, row count, and column info
        """
        import csv

        def extract():
            try:
                # Detect delimiter (CSV vs TSV)
                delimiter = "	" if path.suffix.lower() == ".tsv" else ","

                # Try different encodings
                encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
                rows = None

                for encoding in encodings:
                    try:
                        with open(path, "r", encoding=encoding, newline="") as f:
                            # Sniff to detect actual delimiter if not TSV
                            if path.suffix.lower() != ".tsv":
                                sample = f.read(4096)
                                f.seek(0)
                                try:
                                    dialect = csv.Sniffer().sniff(sample, delimiters=",;	|")
                                    delimiter = dialect.delimiter
                                except csv.Error:
                                    delimiter = ","  # Default to comma

                            reader = csv.reader(f, delimiter=delimiter)
                            rows = list(reader)
                        break
                    except UnicodeDecodeError:
                        continue

                if rows is None:
                    raise Exception("Could not decode CSV file with any supported encoding")

                if not rows:
                    return {
                        "text": "",
                        "pages": 0,
                        "document_metadata": {"columns": 0, "rows": 0},
                    }

                # Format as readable text
                # First row is typically headers
                headers = rows[0] if rows else []
                data_rows = rows[1:] if len(rows) > 1 else []

                text_parts = []

                # Add header line
                if headers:
                    text_parts.append("--- Columns ---")
                    text_parts.append(" | ".join(str(h) for h in headers))
                    text_parts.append("")
                    text_parts.append("--- Data ---")

                # Add data rows
                for row in data_rows:
                    # Format each row, optionally with column names
                    if headers and len(row) == len(headers):
                        # Format as "Column: Value" pairs for better readability
                        pairs = [f"{headers[i]}: {row[i]}" for i in range(len(row))]
                        text_parts.append(" | ".join(pairs))
                    else:
                        text_parts.append(" | ".join(str(cell) for cell in row))

                return {
                    "text": "\n".join(text_parts),
                    "pages": len(data_rows),  # Use row count as page proxy
                    "document_metadata": {
                        "columns": len(headers),
                        "rows": len(data_rows),
                        "headers": headers[:20],  # First 20 column names
                    },
                }

            except Exception as e:
                raise Exception(f"CSV reading error: {str(e)}")

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract)

    # Header-to-metadata-key mapping for email extraction.
    # Entries with tuple values map to multiple metadata keys.
    _EMAIL_HEADER_MAP = {
        "From": ("author", "email_from"),
        "To": "email_to",
        "Cc": "email_cc",
        "Bcc": "email_bcc",
        "Subject": ("title", "email_subject"),
        "Date": ("creation_date", "email_date"),
        "Reply-To": "email_reply_to",
        "Message-ID": "email_message_id",
        "In-Reply-To": "email_in_reply_to",
        "X-Mailer": "creator",
        "User-Agent": "creator",
        "Organization": "organization",
    }

    @staticmethod
    def _strip_emlx_preamble(content: bytes, path: Path) -> bytes:
        """Strip EMLX (Apple Mail) byte-count preamble if present."""
        if path.suffix.lower() != ".emlx":
            return content
        lines = content.split(b"\n", 1)
        if len(lines) > 1:
            try:
                int(lines[0].strip())
                return lines[1]
            except ValueError:
                pass
        return content

    @staticmethod
    def _extract_email_metadata(msg) -> Dict[str, Any]:
        """Extract structured metadata from email message headers."""
        document_metadata: Dict[str, Any] = {}
        for header, keys in ExtractWorker._EMAIL_HEADER_MAP.items():
            value = msg.get(header)
            if not value:
                continue
            if isinstance(keys, tuple):
                for key in keys:
                    document_metadata[key] = str(value)
            else:
                document_metadata[keys] = str(value)
        return document_metadata

    @staticmethod
    def _collect_attachments(msg, document_metadata: Dict[str, Any]) -> None:
        """Scan multipart message for attachments and update metadata."""
        if not msg.is_multipart():
            return
        attachment_count = 0
        attachment_names = []
        for part in msg.walk():
            if part.get_content_disposition() == "attachment":
                attachment_count += 1
                filename = part.get_filename()
                if filename:
                    attachment_names.append(filename)
        if attachment_count > 0:
            document_metadata["attachment_count"] = attachment_count
            if attachment_names:
                document_metadata["attachments"] = ", ".join(attachment_names)

    @staticmethod
    def _extract_email_body(msg) -> tuple:
        """Extract body text parts and part count from email message."""
        import re

        parts = ["--- Body ---"]
        part_count = 0

        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == "text/plain":
                    body = part.get_content()
                    if isinstance(body, str):
                        parts.append(body)
                        part_count += 1
                elif content_type == "text/html" and part_count == 0:
                    html = part.get_content()
                    if isinstance(html, str):
                        text = re.sub(r"<[^>]+>", "", html)
                        text = re.sub(r"\s+", " ", text).strip()
                        if text:
                            parts.append(text)
                            part_count += 1
        else:
            body = msg.get_content()
            if isinstance(body, str):
                parts.append(body)
                part_count = 1
            elif isinstance(body, bytes):
                parts.append(body.decode("utf-8", errors="replace"))
                part_count = 1

        return parts, part_count

    async def _extract_eml(self, path: Path) -> Dict[str, Any]:
        """
        Extract text and metadata from EML or EMLX (Apple Mail) email file.

        Args:
            path: Path to email file

        Returns:
            dict with text (headers + body), part count, and document_metadata
        """
        import email
        from email import policy

        def extract():
            try:
                with open(path, "rb") as f:
                    content = f.read()

                content = self._strip_emlx_preamble(content, path)
                msg = email.message_from_bytes(content, policy=policy.default)

                document_metadata = self._extract_email_metadata(msg)
                self._collect_attachments(msg, document_metadata)

                # Build text header section
                text_parts = []
                header_lines = []
                for header in ["From", "To", "Cc", "Subject", "Date"]:
                    value = msg.get(header)
                    if value:
                        header_lines.append(f"{header}: {value}")
                if header_lines:
                    text_parts.append("--- Headers ---")
                    text_parts.extend(header_lines)
                    text_parts.append("")

                # Extract body
                body_parts, part_count = self._extract_email_body(msg)
                text_parts.extend(body_parts)

                return {
                    "text": "\n".join(text_parts),
                    "pages": max(1, part_count),
                    "document_metadata": document_metadata,
                }

            except Exception as e:
                raise Exception(f"Email file reading error: {str(e)}")

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract)


if __name__ == "__main__":
    """Run the worker if executed directly."""
    from arkham_frame.workers.base import run_worker

    run_worker(ExtractWorker)
